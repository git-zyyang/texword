"""TexWord 后处理器测试。"""

import pytest
from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn

from texword.core.style import StyleConfig
from texword.latex.postprocessor import (
    set_paragraph_font, set_paragraph_spacing, set_page_margins,
    style_headings, style_normal, style_tables, style_captions,
    fix_body_paragraphs, _clean_bib_artifacts, _clean_empty_paragraphs,
    _is_caption_text, _is_bib_entry, _set_widow_orphan, _set_cell_border,
)


@pytest.fixture
def cfg():
    return StyleConfig()


@pytest.fixture
def doc():
    return Document()


class TestHelpers:
    def test_is_caption_figure(self):
        assert _is_caption_text("Figure 1: Results")
        assert _is_caption_text("Fig. 2 shows")
        assert _is_caption_text("Table 3: Summary")
        assert _is_caption_text("表 1 结果")
        assert not _is_caption_text("The figure shows")
        assert not _is_caption_text("")

    def test_is_bib_entry(self):
        assert _is_bib_entry("Acemoglu, D. (2012) Why nations fail")
        assert _is_bib_entry("Zhang, Y. (2024) Digital economy")
        assert not _is_bib_entry("the results show that")
        assert not _is_bib_entry("")


class TestWidowOrphan:
    def test_sets_widow_control(self, doc):
        para = doc.add_paragraph("Test paragraph")
        _set_widow_orphan(para)
        pPr = para._element.pPr
        widow = pPr.find(qn("w:widowControl"))
        assert widow is not None
        assert widow.get(qn("w:val")) == "1"

    def test_no_duplicate_widow_control(self, doc):
        para = doc.add_paragraph("Test")
        _set_widow_orphan(para)
        _set_widow_orphan(para)  # call twice
        pPr = para._element.pPr
        widows = pPr.findall(qn("w:widowControl"))
        assert len(widows) == 1


class TestPageMargins:
    def test_sets_dimensions(self, doc, cfg):
        set_page_margins(doc, cfg)
        section = doc.sections[0]
        assert section.page_width == Cm(cfg.page_width)
        assert section.page_height == Cm(cfg.page_height)
        assert section.top_margin == Cm(cfg.margin_top)


class TestCleanBibArtifacts:
    def test_removes_99(self, doc):
        doc.add_paragraph("99")
        doc.add_paragraph("Real content")
        removed = _clean_bib_artifacts(doc)
        assert removed == 1
        texts = [p.text for p in doc.paragraphs]
        assert "99" not in texts

    def test_keeps_normal_text(self, doc):
        doc.add_paragraph("99 bottles of beer")
        removed = _clean_bib_artifacts(doc)
        assert removed == 0


class TestCleanEmptyParagraphs:
    def test_removes_consecutive_empty(self, doc):
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("")
        doc.add_paragraph("Content")
        removed = _clean_empty_paragraphs(doc)
        assert removed == 2  # keeps first empty, removes 2nd and 3rd

    def test_keeps_single_empty(self, doc):
        doc.add_paragraph("A")
        doc.add_paragraph("")
        doc.add_paragraph("B")
        removed = _clean_empty_paragraphs(doc)
        assert removed == 0


class TestStyleTables:
    def test_three_line_borders(self, doc, cfg):
        table = doc.add_table(rows=3, cols=3)
        table.cell(0, 0).text = "Header"
        table.cell(1, 0).text = "Data"
        style_tables(doc, cfg)

        tblPr = table._tbl.tblPr
        borders = tblPr.find(qn("w:tblBorders"))
        assert borders is not None

        top = borders.find(qn("w:top"))
        assert top.get(qn("w:val")) == "single"
        assert top.get(qn("w:sz")) == "12"

        left = borders.find(qn("w:left"))
        assert left.get(qn("w:val")) == "none"

        insideV = borders.find(qn("w:insideV"))
        assert insideV.get(qn("w:val")) == "none"

    def test_header_bold(self, doc, cfg):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "Header"
        table.cell(1, 0).text = "Data"
        style_tables(doc, cfg)

        for cell in table.rows[0].cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    if run.text.strip():
                        assert run.font.bold is True

    def test_header_cell_bottom_border(self, doc, cfg):
        table = doc.add_table(rows=2, cols=2)
        table.cell(0, 0).text = "H"
        style_tables(doc, cfg)

        cell = table.rows[0].cells[0]
        tcPr = cell._element.find(qn("w:tcPr"))
        tcBorders = tcPr.find(qn("w:tcBorders"))
        bottom = tcBorders.find(qn("w:bottom"))
        assert bottom.get(qn("w:sz")) == "6"


class TestStyleCaptions:
    def test_formats_figure_caption(self, doc, cfg):
        doc.add_paragraph("Figure 1: Test caption")
        count = style_captions(doc, cfg)
        assert count == 1

    def test_skips_non_caption(self, doc, cfg):
        doc.add_paragraph("Normal paragraph text")
        count = style_captions(doc, cfg)
        assert count == 0


class TestSetParagraphFont:
    def test_sets_font(self, doc):
        para = doc.add_paragraph("Hello world")
        set_paragraph_font(para, "Arial", 14, bold=True)
        for run in para.runs:
            assert run.font.name == "Arial"
            assert run.font.size == Pt(14)
            assert run.font.bold is True
