"""python-docx 后处理：修复字体、行距、标题、表格、参考文献等格式。"""

import re

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from texword.core.style import StyleConfig


# ── 段落工具函数 ──

def set_paragraph_font(para, font_name: str, font_size_pt: float,
                       bold=False, italic=False, color=None):
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(font_size_pt)
        run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        if bold:
            run.font.bold = True
        if italic:
            run.font.italic = True
        if color:
            run.font.color.rgb = color


def set_paragraph_spacing(para, line_spacing: float = 2.0,
                          space_before: float = 0, space_after: float = 0):
    pf = para.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = line_spacing
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)


# ── 页面设置 ──

def set_page_margins(doc, cfg: StyleConfig):
    for section in doc.sections:
        section.page_width = Cm(cfg.page_width)
        section.page_height = Cm(cfg.page_height)
        section.top_margin = Cm(cfg.margin_top)
        section.bottom_margin = Cm(cfg.margin_bottom)
        section.left_margin = Cm(cfg.margin_left)
        section.right_margin = Cm(cfg.margin_right)


# ── 样式设置 ──

def style_headings(doc, cfg: StyleConfig):
    heading_sizes = {
        "Heading 1": cfg.font_size_h1,
        "Heading 2": cfg.font_size_h2,
        "Heading 3": cfg.font_size_h3,
    }
    for style_name, size in heading_sizes.items():
        try:
            style = doc.styles[style_name]
            style.font.name = cfg.font_body
            style.font.size = Pt(size)
            style.font.bold = True
            style.font.color.rgb = RGBColor(0, 0, 0)
            style.element.rPr.rFonts.set(qn("w:eastAsia"), cfg.font_cjk)
            pf = style.paragraph_format
            pf.space_before = Pt(12)
            pf.space_after = Pt(6)
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            pf.line_spacing = cfg.line_spacing
        except KeyError:
            pass


def style_normal(doc, cfg: StyleConfig):
    try:
        style = doc.styles["Normal"]
        style.font.name = cfg.font_body
        style.font.size = Pt(cfg.font_size_body)
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.element.rPr.rFonts.set(qn("w:eastAsia"), cfg.font_cjk)
        pf = style.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = cfg.line_spacing
        pf.first_line_indent = Cm(cfg.first_line_indent)
    except KeyError:
        pass


def _set_cell_border(cell, edge, sz="4", val="single", color="000000"):
    """设置单元格某条边框。"""
    tcPr = cell._element.get_or_add_tcPr()
    borders = tcPr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tcPr.append(borders)
    # 移除已有的同名边框
    for existing in borders.findall(qn(f"w:{edge}")):
        borders.remove(existing)
    el = OxmlElement(f"w:{edge}")
    el.set(qn("w:val"), val)
    el.set(qn("w:sz"), sz)
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), color)
    borders.append(el)


def style_tables(doc, cfg: StyleConfig):
    for table in doc.tables:
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        rows = table.rows
        for i, row in enumerate(rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in para.runs:
                        run.font.name = cfg.font_body
                        run.font.size = Pt(cfg.font_size_table)
                        run.font.element.rPr.rFonts.set(
                            qn("w:eastAsia"), cfg.font_cjk)
                        # 表头行加粗
                        if i == 0:
                            run.font.bold = True
                    pf = para.paragraph_format
                    pf.line_spacing_rule = WD_LINE_SPACING.SINGLE
                    pf.space_before = Pt(1)
                    pf.space_after = Pt(1)
                    pf.first_line_indent = Cm(0)

        # 三线表边框（学术 booktabs 风格）
        tbl = table._tbl
        tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement("w:tblPr")
        # 移除已有的 tblBorders
        for existing in tblPr.findall(qn("w:tblBorders")):
            tblPr.remove(existing)
        borders = OxmlElement("w:tblBorders")
        # 顶线（粗）
        for edge, sz, val in [
            ("top", "12", "single"),      # 1.5pt 粗顶线
            ("bottom", "12", "single"),    # 1.5pt 粗底线
            ("left", "0", "none"),
            ("right", "0", "none"),
            ("insideH", "0", "none"),
            ("insideV", "0", "none"),
        ]:
            el = OxmlElement(f"w:{edge}")
            el.set(qn("w:val"), val)
            el.set(qn("w:sz"), sz)
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tblPr.append(borders)

        # 表头行底部加中等粗线（分隔表头和数据）
        if rows:
            for cell in rows[0].cells:
                _set_cell_border(cell, "bottom", sz="6", val="single")


# ── 辅助函数 ──

def _set_widow_orphan(para):
    """设置寡行孤行控制。"""
    pPr = para._element.get_or_add_pPr()
    # 移除已有的 widowControl
    for existing in pPr.findall(qn("w:widowControl")):
        pPr.remove(existing)
    widow = OxmlElement("w:widowControl")
    widow.set(qn("w:val"), "1")
    pPr.insert(0, widow)


def _is_caption_text(text):
    """判断是否为图表标题文本。"""
    return bool(re.match(
        r"^(Figure|Fig\.|Table|图|表)\s*\d", text, re.IGNORECASE))


def _is_bib_entry(text):
    """判断是否为参考文献条目（Author (Year) 模式）。"""
    return bool(re.match(r"^[A-Z].*\(\d{4}", text))


# ── 清理 thebibliography 残留 ──

def _clean_bib_artifacts(doc):
    """移除 pandoc 转换 thebibliography{99} 产生的残留文本。"""
    removed = 0
    for para in list(doc.paragraphs):
        text = para.text.strip()
        # pandoc 会把 \begin{thebibliography}{99} 的 {99} 转为独立段落
        if text in ("99", "100") and len(para.runs) <= 1:
            parent = para._element.getparent()
            parent.remove(para._element)
            removed += 1
    return removed


# ── 清理多余空段落 ──

def _clean_empty_paragraphs(doc):
    """移除连续的空段落（保留单个空段落作为分隔）。"""
    removed = 0
    prev_empty = False
    for para in list(doc.paragraphs):
        is_empty = not para.text.strip() and not para._element.findall(
            f".//{{{qn('w:drawing')}}}")
        if is_empty and prev_empty:
            parent = para._element.getparent()
            parent.remove(para._element)
            removed += 1
        else:
            prev_empty = is_empty
    return removed


# ── 图表标题格式化 ──

def style_captions(doc, cfg: StyleConfig):
    """格式化图表标题（Figure X / Table X 开头的段落）。"""
    count = 0
    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        # pandoc 生成的 caption 通常是 "Image Caption" 或 "Caption" 样式
        # 也可能是普通段落但以 Figure/Table 开头
        is_caption = (style_name in ("Image Caption", "Caption", "Table Caption")
                      or _is_caption_text(text))

        if is_caption and text:
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = cfg.font_body
                run.font.size = Pt(cfg.font_size_caption)
                run.font.element.rPr.rFonts.set(qn("w:eastAsia"), cfg.font_cjk)
            set_paragraph_spacing(para, 1.5, 6, 6)
            para.paragraph_format.first_line_indent = Cm(0)
            count += 1
    return count


# ── 正文段落修复 ──

def fix_body_paragraphs(doc, cfg: StyleConfig):
    in_references = False
    found_conclusion = False

    for para in doc.paragraphs:
        style_name = para.style.name if para.style else ""
        text = para.text.strip()

        # 检测参考文献区域
        if style_name.startswith("Heading"):
            if "conclusion" in text.lower():
                found_conclusion = True
                in_references = False
            elif "reference" in text.lower() or "bibliography" in text.lower():
                in_references = True
            else:
                in_references = False
        if (found_conclusion and style_name == "Body Text" and not in_references
                and _is_bib_entry(text)):
            in_references = True

        # 跳过 caption（已单独处理）
        if _is_caption_text(text) or style_name in (
                "Image Caption", "Caption", "Table Caption"):
            continue

        # Title
        if style_name == "Title":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = cfg.font_body
                run.font.size = Pt(cfg.font_size_title)
                run.font.bold = True
            set_paragraph_spacing(para, cfg.line_spacing, 24, 12)
            _set_widow_orphan(para)
            continue

        # Author / Date
        if style_name in ("Author", "Date"):
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = cfg.font_body
                run.font.size = Pt(cfg.font_size_body)
            set_paragraph_spacing(para, 1.5, 6, 6)
            para.paragraph_format.first_line_indent = Cm(0)
            continue

        # Abstract Title
        if style_name == "Abstract Title":
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in para.runs:
                run.font.name = cfg.font_body
                run.font.size = Pt(cfg.font_size_h2)
                run.font.bold = True
            set_paragraph_spacing(para, cfg.line_spacing, 12, 6)
            para.paragraph_format.first_line_indent = Cm(0)
            continue

        # Abstract
        if style_name == "Abstract":
            for run in para.runs:
                run.font.name = cfg.font_body
                run.font.size = Pt(cfg.font_size_abstract)
            set_paragraph_spacing(para, 1.5, 0, 0)
            para.paragraph_format.first_line_indent = Cm(0)
            para.paragraph_format.left_indent = Cm(1.27)
            para.paragraph_format.right_indent = Cm(1.27)
            continue

        # Headings — 显式设置间距（不依赖样式继承）
        if style_name.startswith("Heading"):
            para.paragraph_format.first_line_indent = Cm(0)
            level = style_name[-1] if style_name[-1].isdigit() else "1"
            sizes = {"1": cfg.font_size_h1, "2": cfg.font_size_h2,
                     "3": cfg.font_size_h3}
            sz = sizes.get(level, cfg.font_size_h1)
            for run in para.runs:
                run.font.name = cfg.font_body
                run.font.size = Pt(sz)
                run.font.bold = True
                run.font.color.rgb = RGBColor(0, 0, 0)
                run.font.element.rPr.rFonts.set(qn("w:eastAsia"), cfg.font_cjk)
            set_paragraph_spacing(para, cfg.line_spacing, 12, 6)
            _set_widow_orphan(para)
            continue

        # References (APA style: hanging indent, double spacing, preserve italic)
        if in_references and style_name == "Body Text":
            for run in para.runs:
                run.font.name = cfg.font_body
                run.font.size = Pt(cfg.font_size_ref)
                # 保留 pandoc 转换的斜体（期刊名、书名等）
            set_paragraph_spacing(para, cfg.line_spacing, 0, 0)
            para.paragraph_format.first_line_indent = Cm(-1.27)
            para.paragraph_format.left_indent = Cm(1.27)
            continue

        # Normal paragraphs — force correct font/size
        for run in para.runs:
            run.font.name = cfg.font_body
            if not run.font.size or run.font.size != Pt(cfg.font_size_body):
                run.font.size = Pt(cfg.font_size_body)
            try:
                run.font.element.rPr.rFonts.set(
                    qn("w:eastAsia"), cfg.font_cjk)
            except Exception:
                pass

        pf = para.paragraph_format
        pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
        pf.line_spacing = cfg.line_spacing

        # 所有正文段落添加寡行孤行控制
        if text:
            _set_widow_orphan(para)


# ── 公式编号 ──

MATH_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"


def number_display_equations(doc, cfg: StyleConfig):
    """为显示公式（oMathPara）添加右对齐编号 (1), (2), ...

    使用 Word 标准方式：右对齐制表位 + 编号文本。
    """
    page_width_emu = int(Cm(cfg.page_width).emu)
    margin_emu = int(Cm(cfg.margin_left).emu + Cm(cfg.margin_right).emu)
    right_pos = page_width_emu - margin_emu  # 文本区域右边界

    eq_num = 0
    for para in doc.paragraphs:
        omath_paras = para._element.findall(
            f".//{{{MATH_NS}}}oMathPara")
        if not omath_paras:
            continue

        eq_num += 1

        # 添加右对齐制表位
        pPr = para._element.get_or_add_pPr()
        tabs = pPr.find(qn("w:tabs"))
        if tabs is None:
            tabs = OxmlElement("w:tabs")
            pPr.append(tabs)
        tab = OxmlElement("w:tab")
        tab.set(qn("w:val"), "right")
        tab.set(qn("w:pos"), str(right_pos // 635))  # EMU → twips
        tab.set(qn("w:leader"), "none")
        tabs.append(tab)

        # 添加 tab + 编号 run
        run_el = OxmlElement("w:r")
        rPr = OxmlElement("w:rPr")
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), cfg.font_body)
        rFonts.set(qn("w:hAnsi"), cfg.font_body)
        rPr.append(rFonts)
        sz = OxmlElement("w:sz")
        sz.set(qn("w:val"), str(cfg.font_size_body * 2))  # half-points
        rPr.append(sz)
        run_el.append(rPr)

        tab_char = OxmlElement("w:tab")
        run_el.append(tab_char)

        t = OxmlElement("w:t")
        t.set(qn("xml:space"), "preserve")
        t.text = f"({eq_num})"
        run_el.append(t)

        para._element.append(run_el)

    return eq_num


# ── 页眉页脚 ──

def add_header_footer(doc, title_short: str = ""):
    """去掉页眉，页码居中放在页脚。"""
    for section in doc.sections:
        # ── 清空页眉 ──
        header = section.header
        header.is_linked_to_previous = False
        for para in header.paragraphs:
            para.text = ""
            for run in list(para.runs):
                run._r.getparent().remove(run._r)

        # ── 页脚：居中页码 ──
        footer = section.footer
        footer.is_linked_to_previous = False
        if not footer.paragraphs:
            footer.add_paragraph()
        fp = footer.paragraphs[0]
        fp.text = ""
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        run = fp.add_run()
        run.font.size = Pt(10)
        run.font.name = "Times New Roman"

        fld_char_begin = OxmlElement("w:fldChar")
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        run._r.append(fld_char_begin)

        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = " PAGE "
        run._r.append(instr)

        fld_char_sep = OxmlElement("w:fldChar")
        fld_char_sep.set(qn("w:fldCharType"), "separate")
        run._r.append(fld_char_sep)

        fld_char_end = OxmlElement("w:fldChar")
        fld_char_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_char_end)


# ── 主入口 ──

def postprocess(docx_path: str, output_path: str, cfg: StyleConfig,
                title_short: str = ""):
    """后处理 pandoc 输出的 docx，修复格式。"""
    print("\n[Phase 3] 后处理格式...")

    doc = Document(docx_path)

    print("  设置页面尺寸和边距...")
    set_page_margins(doc, cfg)

    print("  清理转换残留...")
    n_artifacts = _clean_bib_artifacts(doc)
    n_empty = _clean_empty_paragraphs(doc)
    if n_artifacts:
        print(f"    移除 {n_artifacts} 个残留文本")
    if n_empty:
        print(f"    移除 {n_empty} 个多余空段落")

    print("  设置标题样式...")
    style_headings(doc, cfg)
    style_normal(doc, cfg)

    print("  修复正文字体和行距...")
    fix_body_paragraphs(doc, cfg)

    print("  格式化图表标题...")
    n_captions = style_captions(doc, cfg)
    if n_captions:
        print(f"    处理 {n_captions} 个图表标题")

    print("  美化表格...")
    style_tables(doc, cfg)

    print("  添加公式编号...")
    n_eq = number_display_equations(doc, cfg)
    if n_eq:
        print(f"    编号 {n_eq} 个显示公式")

    print("  设置页脚页码...")
    add_header_footer(doc, title_short)

    doc.save(output_path)
    import os
    size_kb = os.path.getsize(output_path) / 1024
    print(f"\n  最终输出 → {output_path} ({size_kb:.1f} KB)")
