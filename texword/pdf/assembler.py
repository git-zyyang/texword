"""PDF→DOCX 组装器 — 将提取的内容块组装为 Word 文档。"""

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from texword.core.style import StyleConfig
from texword.pdf.extractor import ContentBlock, BlockType
from texword.latex.postprocessor import (
    set_paragraph_spacing, style_headings, style_normal,
    style_tables, set_page_margins
)


class PDFAssembler:
    """将 ContentBlock 列表组装为格式化的 Word 文档。"""

    def __init__(self, cfg: StyleConfig = None):
        self.cfg = cfg or StyleConfig()

    def assemble(self, blocks: list[ContentBlock], output_path: str):
        """将内容块组装为 docx。"""
        doc = Document()

        # 页面设置
        set_page_margins(doc, self.cfg)
        style_headings(doc, self.cfg)
        style_normal(doc, self.cfg)

        for block in blocks:
            if block.type == BlockType.TITLE:
                p = doc.add_heading(block.text, level=0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif block.type == BlockType.AUTHOR:
                p = doc.add_paragraph(block.text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.style = doc.styles.get("Author", doc.styles["Normal"])

            elif block.type == BlockType.ABSTRACT:
                doc.add_heading("Abstract", level=1)
                p = doc.add_paragraph(block.text)
                set_paragraph_spacing(p, 1.5)

            elif block.type == BlockType.HEADING:
                doc.add_heading(block.text, level=1)

            elif block.type == BlockType.FIGURE:
                if block.image_path:
                    doc.add_picture(block.image_path, width=Inches(5))

            elif block.type == BlockType.EQUATION:
                # 如果有 LaTeX，后续可转 OMML
                p = doc.add_paragraph(block.latex or block.text)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER

            elif block.type == BlockType.REFERENCE:
                p = doc.add_paragraph(block.text)
                pf = p.paragraph_format
                pf.first_line_indent = Cm(-1.27)
                pf.left_indent = Cm(1.27)
                set_paragraph_spacing(p, 1.5, 0, 3)

            else:  # TEXT
                doc.add_paragraph(block.text)

        # 美化表格
        style_tables(doc, self.cfg)

        doc.save(output_path)
        print(f"  PDF→DOCX 组装完成 → {output_path}")
